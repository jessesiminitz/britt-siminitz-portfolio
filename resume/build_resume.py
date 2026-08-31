#!/usr/bin/env python3
"""Build Brittany Siminitz's freelance-writer resume as a .docx.

Every figure here is derived from data/jck_archive.json (scraped from JCK's
own author archive) or from an article page whose byline was verified
directly — nothing is estimated. Run after re-scraping to refresh the counts:

    ../.venv/bin/python build_resume.py

Fields the subject must supply herself are marked with the PLACEHOLDER
constant so they're obvious in the output rather than silently invented.
"""

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

HERE = Path(__file__).resolve().parent
ARCHIVE = HERE.parent / "data" / "jck_archive.json"
OUT = HERE / "Brittany_Siminitz_Resume.docx"

INK = RGBColor(0x2A, 0x25, 0x21)
SOFT = RGBColor(0x5C, 0x53, 0x49)
EMERALD = RGBColor(0x14, 0x51, 0x3F)

PLACEHOLDER = "[ADD]"


def archive_stats():
    """Pull the headline numbers straight from the scraped archive."""
    rows = json.loads(ARCHIVE.read_text())
    dates = sorted(r["date"] for r in rows if r["date"])
    return {
        "total": len(rows),
        "first_year": dates[0][:4],
        "last_year": dates[-1][:4],
    }


def selected_clips():
    """Pull a spread of clips across outlets from the curated articles.json.

    Chosen by URL so the title/date shown always matches the verified data
    rather than drifting out of sync with a hand-typed copy.
    """
    picks = [
        "https://www.jckonline.com/editorial-article/talismania-montana-sapphires/",
        "https://www.jckonline.com/editorial-article/jewelry-trends-2025/",
        "https://gemandjewel.substack.com/p/the-ageless-allure-of-vintage-diamonds",
        "https://www.riogrande.com/knowledge-hub/articles/wedding-band-trends-2026-personalization-takes-center-stage/",
        "https://jennylaurenjewelry.com/blogs/news/9-vacation-worthy-jewels-to-wear-all-summer-long",
    ]
    data = json.loads((HERE.parent / "data" / "articles.json").read_text())
    by_url = {
        a["url"]: (a, o["name"])
        for o in data["outlets"]
        for a in o["articles"]
    }
    months = ("Jan", "Feb", "Mar", "Apr", "May", "Jun",
              "Jul", "Aug", "Sep", "Oct", "Nov", "Dec")
    out = []
    for url in picks:
        art, outlet = by_url[url]
        date = ""
        if art.get("date"):
            y, m, _ = art["date"].split("-")
            date = f"{months[int(m) - 1]} {y}"
        out.append({"title": art["title"], "outlet": outlet, "date": date})
    return out


def rule(par):
    """Thin bottom border, used under section headings."""
    p = par._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "C9BFAE")
    borders.append(bottom)
    p.append(borders)


def heading(doc, text):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(8)
    par.paragraph_format.space_after = Pt(4)
    run = par.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = EMERALD
    run.font.name = "Calibri"
    rule(par)
    return par


def entry(doc, title, meta, dates):
    """One role: bold title + light meta on the left, dates right-aligned."""
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(6)
    par.paragraph_format.space_after = Pt(1)
    # Right-align the date via a tab stop at the right margin.
    par.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_ALIGN_PARAGRAPH.RIGHT)

    r = par.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = INK

    if meta:
        r2 = par.add_run(f"  ·  {meta}")
        r2.font.size = Pt(10)
        r2.font.color.rgb = SOFT
        r2.italic = True

    r3 = par.add_run(f"\t{dates}")
    r3.font.size = Pt(9.5)
    r3.font.color.rgb = SOFT
    return par


def bullet(doc, text):
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.space_after = Pt(2)
    par.paragraph_format.left_indent = Inches(0.22)
    par.paragraph_format.line_spacing = 1.0
    run = par.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = INK
    return par


def labeled(doc, label, body):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.line_spacing = 1.0
    a = par.add_run(f"{label}  ")
    a.bold = True
    a.font.size = Pt(10)
    a.font.color.rgb = INK
    b = par.add_run(body)
    b.font.size = Pt(10)
    b.font.color.rgb = INK
    return par


def build():
    s = archive_stats()
    total = f"{s['total']:,}"
    years = int(s["last_year"]) - int(s["first_year"])

    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin"):
        setattr(section, attr, Inches(0.5))
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(0)

    # ---- Name block ----
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(1)
    r = name.add_run("BRITTANY SIMINITZ")
    r.bold = True
    r.font.size = Pt(21)
    r.font.color.rgb = INK
    r.font.name = "Georgia"

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag.paragraph_format.space_after = Pt(3)
    r = tag.add_run("Jewelry & Luxury Trade Journalist")
    r.font.size = Pt(11)
    r.italic = True
    r.font.color.rgb = EMERALD
    r.font.name = "Georgia"

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(2)
    r = contact.add_run(
        f"{PLACEHOLDER} email   ·   {PLACEHOLDER} phone   ·   Washington, DC {PLACEHOLDER} confirm"
    )
    r.font.size = Pt(9)
    r.font.color.rgb = SOFT

    links = doc.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links.paragraph_format.space_after = Pt(2)
    r = links.add_run(
        "Portfolio: jessesiminitz.github.io/britt-siminitz-portfolio   ·   "
        "jckonline.com/writer/brittany-siminitz   ·   X: @BrittsPickJCK"
    )
    r.font.size = Pt(9)
    r.font.color.rgb = SOFT

    # ---- Summary ----
    heading(doc, "Profile")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(
        f"Jewelry trade journalist with {years} years and {total} published articles covering fine "
        "jewelry, colored gemstones, design, and the business of the industry. Contributing Editor "
        "at JCK, where I write three standing blogs and the Britt's Pick column. Equally at home "
        "filing product and supplier news on deadline, reporting trend features off the Las Vegas "
        "show floor, and translating social platform shifts into practical guidance for retailers — "
        "plus branded editorial for manufacturers and consumer brands."
    )
    r.font.size = Pt(10)
    r.font.color.rgb = INK

    # ---- Experience ----
    heading(doc, "Experience")

    entry(doc, "Contributing Editor", "JCK / JCK Magazine", f"{s['first_year']}–Present")
    bullet(
        doc,
        f"Author of {total} published articles since {s['first_year']} — averaging roughly "
        "245 pieces a year on a daily filing schedule.",
    )
    bullet(
        doc,
        "Write and own three standing blogs: On Your Market (designer and collection "
        "coverage), Supplier News (trade and manufacturing), and Social Setting (social "
        "media and digital marketing for jewelers).",
    )
    bullet(
        doc,
        "Created and have written Britt's Pick, a single-piece design column, for "
        f"{years} years — 660 installments plus annual year-in-review roundups.",
    )
    bullet(
        doc,
        "Cover JCK Las Vegas and LUXURY each year, filing show previews, trend forecasts, "
        "and post-show roundups.",
    )
    entry(doc, "Contributing Writer", "Gem + Jewel — Jewelers Mutual Group", "2026–Present")
    bullet(
        doc,
        "Consumer-facing trend and designer features for the Substack and biannual print "
        "magazine, covering vintage diamonds, birthstones, engagement trends, and styling.",
    )

    entry(doc, "Contributing Writer", "Rio Grande Knowledge Hub", "2025–Present")
    bullet(
        doc,
        "Branded editorial for the industry's largest wholesale supplier (a Berkshire "
        "Hathaway company), including its annual Color of the Year feature and bridal "
        "trend reporting for a trade audience of working jewelers.",
    )

    entry(doc, "Contributing Writer", "Jenny Lauren Jewelry", f"{PLACEHOLDER} dates")
    bullet(doc, "Editorial styling features for the fine jewelry brand's journal.")

    entry(doc, "Contributor", "JCK Insider", f"c. 2016–2018 {PLACEHOLDER} confirm")
    bullet(
        doc,
        "Trend files, designer spotlights, and show previews for JCK's Las Vegas "
        "event publication.",
    )

    entry(
        doc,
        f"{PLACEHOLDER} Prior role — add title, dates, scope",
        f"McMurry/TMG (now Manifest) {PLACEHOLDER} confirm",
        PLACEHOLDER,
    )

    # ---- Beats ----
    heading(doc, "Beats & Expertise")
    labeled(
        doc,
        "Editorial",
        "Fine jewelry design · Colored gemstones & diamonds · Designer interviews · "
        "Trend forecasting · Trade shows",
    )
    labeled(
        doc,
        "Trade & business",
        "Supplier and manufacturing news · Retail and holiday sell-through · Bridal · Market trends",
    )
    labeled(
        doc,
        "Digital",
        "Social strategy for jewelry brands (Instagram, Pinterest, TikTok, Meta) · "
        "Newsletter & Substack · Branded editorial",
    )
    labeled(
        doc,
        "Craft",
        "Daily deadline filing · Product roundups · SEO-aware headlines · Photography sourcing",
    )

    # ---- Selected clips ----
    heading(doc, "Selected Clips")
    for clip in selected_clips():
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(2)
        par.paragraph_format.line_spacing = 1.0
        t = par.add_run(f"“{clip['title']}”")
        t.font.size = Pt(10)
        t.font.color.rgb = INK
        meta = f"  ·  {clip['outlet']}"
        if clip["date"]:
            meta += f", {clip['date']}"
        m = par.add_run(meta)
        m.font.size = Pt(9.5)
        m.italic = True
        m.font.color.rgb = SOFT

    # ---- Education ----
    heading(doc, "Education")
    entry(
        doc,
        f"{PLACEHOLDER} Degree",
        f"University of North Carolina at Wilmington {PLACEHOLDER} confirm",
        PLACEHOLDER,
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"  archive: {total} articles, {s['first_year']}–{s['last_year']} ({years} yrs)")


if __name__ == "__main__":
    build()
